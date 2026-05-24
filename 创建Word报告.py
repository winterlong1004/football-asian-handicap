from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 设置文档默认字体
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# 添加页眉
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = "足球博彩分析系统 - Rule优化报告"
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加页脚
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "第 "
run = footer_para.add_run()
run._element.append(docx.oxml.parse_xml('<w:fldChar w:fldCharType="begin"/>'))
run = footer_para.add_run()
run._element.append(docx.oxml.parse_xml('<w:instrText xml:space="preserve"> PAGE </w:instrText>'))
run = footer_para.add_run()
run._element.append(docx.oxml.parse_xml('<w:fldChar w:fldCharType="end"/>'))
footer_para.add_run(" 页，分析日期：2026-05-10")
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 标题
title = doc.add_heading('足球博彩盘口分析Rule系统优化', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)

p = doc.add_paragraph()
p.add_run('基于44场复盘数据的深度分析，以下是核心发现和优化建议：').bold = True

# 关键指标表格
table1 = doc.add_table(rows=4, cols=3)
table1.style = 'Light Grid Accent 1'
table1.rows[0].cells[0].text = '指标'
table1.rows[0].cells[1].text = '数值'
table1.rows[0].cells[2].text = '评价'

data1 = [
    ('双选命中率', '82.5% (33/40)', '⭐⭐⭐⭐⭐ 优秀'),
    ('单选命中率', '34.3% (12/35)', '⭐⭐ 需大幅提升'),
    ('临场盘单选命中率', '44.4% (4/9)', '⭐⭐⭐ 有改善空间'),
    ('Rule 13命中率', '100% (3/3)', '⭐⭐⭐⭐⭐ 极强信号')
]

for i, (metric, value, eval) in enumerate(data1, 1):
    table1.rows[i].cells[0].text = metric
    table1.rows[i].cells[1].text = value
    table1.rows[i].cells[2].text = eval

doc.add_paragraph()

# 二、Rule有效性验证
doc.add_heading('二、Rule有效性验证（基于44场复盘）', 1)

doc.add_heading('1. 高效规则（保留并优先使用）', 2)
table2 = doc.add_table(rows=5, cols=6)
table2.style = 'Light Grid Accent 2'
headers2 = ['规则', '名称', '样本数', '单选命中', '双选命中', '结论']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h

data2 = [
    ('Rule 13', '欧赔超级一致', '3场', '3/3 (100%)', '3/3 (100%)', '✅ 黄金规则'),
    ('Rule 6', '明*主水≥1.10', '1场', '1/1 (100%)', '1/1 (100%)', '✅ 保留（需更多样本）'),
    ('Rule 22', '主水大幅降(≥5点)', '3场', '2/3 (67%)', '2/3 (67%)', '✅ 保留继续验证')
]

for i, row_data in enumerate(data2, 1):
    for j, value in enumerate(row_data):
        table2.rows[i].cells[j].text = value

doc.add_paragraph()

doc.add_heading('2. 需优化的规则（双选命中率偏低）', 2)
table3 = doc.add_table(rows=4, cols=5)
table3.style = 'Light Grid Accent 3'
headers3 = ['规则', '名称', '样本数', '问题', '优化建议']
for i, h in enumerate(headers3):
    table3.rows[0].cells[i].text = h

data3 = [
    ('Rule 19', '浅盘禁止单选', '4场', '双选仅50% (2/4)', '浅盘双选策略需调整'),
    ('Rule 21', '过度一致性预警', '2场', '双选0% (0/2)', '⚠️ 严重问题！需修正'),
    ('Rule 21详情', 'R40,R41失败', '2场', '过度一致+双选失败', '应反向投注或观望')
]

for i, row_data in enumerate(data3, 1):
    for j, value in enumerate(row_data):
        table3.rows[i].cells[j].text = value

doc.add_paragraph()

doc.add_heading('3. 待验证规则（样本不足）', 2)
table4 = doc.add_table(rows=5, cols=4)
table4.style = 'Light Grid Accent 4'
headers4 = ['观察', '描述', '样本数', '状态']
for i, h in enumerate(headers4):
    table4.rows[0].cells[i].text = h

data4 = [
    ('观察D', '欧赔诱盘(一升一降)', '2场(R43,R44)', '需1例升级为正式Rule'),
    ('观察A', '香港马*逆势低水≤0.85', '1场', '需≥2例验证'),
    ('Rule 16', '深盘主水全破1.0', '1场', '需≥3例验证'),
    ('Rule 17', '隐性大热倒灶', '1场', '需≥3例验证')
]

for i, row_data in enumerate(data4, 1):
    for j, value in enumerate(row_data):
        table4.rows[i].cells[j].text = value

doc.add_paragraph()

# 三、关键优化建议
doc.add_heading('三、关键优化建议', 1)

doc.add_heading('建议1：Rule 21 过度一致性预警策略修正（最高优先级）', 2)
p = doc.add_paragraph()
p.add_run('当前问题：').bold = True
p.add_run('Rule 21触发时，双选也失败（0/2 = 0%）')

p = doc.add_paragraph()
p.add_run('修正方案：').bold = True
p.add_run('\n原方案：过度一致 → 禁止单选 + 双选"不败"')
p.add_run('\n修正方案A：过度一致(浅盘) → 观望 或 反向投注（小注）')
p.add_run('\n修正方案B：过度一致(深盘) → "分胜负"而非"不败"')

doc.add_paragraph()

doc.add_heading('建议2：Rule 19 浅盘双选策略优化（高优先级）', 2)
p = doc.add_paragraph()
p.add_run('关键提醒：').bold = True
p.add_run('Rule 19只禁止单选，不意味着必须投双选。浅盘+亚盘分歧 = 可以观望！')

doc.add_paragraph()

doc.add_heading('建议3：观察D升级为正式Rule（中优先级）', 2)
p = doc.add_paragraph()
p.add_run('升级条件：').bold = True
p.add_run('还需至少1例验证才能升级为正式Rule。当前可标记为准正式规则（需1例）。')

doc.add_paragraph()

doc.add_heading('建议4：单选策略根本性改革（战略级）', 2)
p = doc.add_paragraph()
p.add_run('改革方案：').bold = True)
p.add_run('\n1. 所有单选必须基于"临场最终盘"，禁止早期盘单选')
p.add_run('\n2. 单选前置条件更严格：必须同时满足 Rule 13 OR Rule 22，且无 Rule 19/21 冲突')
p.add_run('\n3. 其他情况一律用双选或观望')

doc.add_paragraph()

# 四、优化后的Rule优先级体系
doc.add_heading('四、优化后的Rule优先级体系', 1)

table5 = doc.add_table(rows=8, cols=5)
table5.style = 'Light Grid Accent 5'
headers5 = ['优先级', '规则名称', '触发条件', '应对策略', '预期命中率']
for i, h in enumerate(headers5):
    table5.rows[0].cells[i].text = h

data5 = [
    ('P0', 'Rule 13', '欧赔4家超级一致', '高信心单选(5星)', '~100%'),
    ('P1', 'Rule 6', '明*主水≥1.10', '看衰该队(反向)', '~100%'),
    ('P1', 'Rule 22', '主水降≥5点', '看好主队', '~67%'),
    ('P2', '观察D(准)', '欧赔一升一降+亚盘多数', '信亚盘多数', '待验证'),
    ('P3', 'Rule 19', '浅盘(0/0.5)', '禁止单选，谨慎双选', '~50%'),
    ('P3', 'Rule 21(修正)', '过度一致性预警', '反向投注或观望', '待验证'),
    ('P4', 'Rule 16/17', '深盘特殊场景', '降低信心度', '不足')
]

for i, row_data in enumerate(data5, 1):
    for j, value in enumerate(row_data):
        table5.rows[i].cells[j].text = value

doc.add_paragraph()

# 五、具体行动清单
doc.add_heading('五、具体行动清单', 1)

doc.add_heading('立即执行（下次比赛即可应用）', 2)
doc.add_paragraph('✅ Rule 21修正：过度一致时考虑反向或观望')
doc.add_paragraph('✅ Rule 19优化：浅盘+亚盘分歧=可以观望')
doc.add_paragraph('✅ 单选限制：只基于临场最终盘')

doc.add_heading('短期目标（积累10场新数据后）', 2)
doc.add_paragraph('⬜ 观察D升级：还需1例即可升级为正式Rule')
doc.add_paragraph('⬜ Rule 22验证：再积累2-3例确认67%命中率稳定性')
doc.add_paragraph('⬜ Rule 21修正验证：测试新策略是否有效')

doc.add_heading('中期目标（积累20场新数据后）', 2)
doc.add_paragraph('⬜ Rule 16/17验证：积累足够样本确认有效性')
doc.add_paragraph('⬜ 观察A验证：香港马*逆势低水规律确认')
doc.add_paragraph('⬜ 联赛差异性分析：瑞典超、韩K联、五大联赛是否有不同规律')

doc.add_paragraph()

# 六、预期改进效果
doc.add_heading('六、预期改进效果', 1)

table6 = doc.add_table(rows=5, cols=4)
table6.style = 'Light Grid Accent 6'
headers6 = ['指标', '当前值', '目标值', '改进幅度']
for i, h in enumerate(headers6):
    table6.rows[0].cells[i].text = h

data6 = [
    ('双选命中率', '82.5%', '88%+', '+5.5%'),
    ('单选命中率', '34.3%', '50%+', '+15.7%'),
    ('临场盘单选', '44.4%', '60%+', '+15.6%'),
    ('Rule 21场景', '0%', '50%+', '+50%')
]

for i, row_data in enumerate(data6, 1):
    for j, value in enumerate(row_data):
        table6.rows[i].cells[j].text = value

doc.add_paragraph()

# 七、总结
doc.add_heading('七、总结', 1)

p = doc.add_paragraph()
p.add_run('核心结论：').bold = True)
p.add_run('\n1. Rule 13（欧赔超级一致）是当前最强信号，应给予最高权重')
p.add_run('\n2. Rule 21需要重大修正：过度一致时应反向或观望，而非盲目跟风')
p.add_run('\n3. 临场最终盘远优于早期盘口，所有决策应基于最终盘')
p.add_run('\n4. 单选必须极其严格：只有极强信号才允许单选')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('⚠️ 风险提示：').bold = True
p.add_run('\n- 总样本量44场仍偏少，统计显著性有限')
p.add_run('\n- 新规则需要至少3例验证才能确立')
p.add_run('\n- 不同联赛可能有不同规律（需分类验证）')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('📝 下一步行动：').bold = True
p.add_run('\n1. 下次比赛立即应用修正后的Rule 21策略')
p.add_run('\n2. 继续积累复盘数据，重点关注观察D、Rule 22的新案例')
p.add_run('\n3. 每10场数据进行一次Rule系统回顾更新')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('分析完成时间：2026-05-10 23:50').italic = True
p.add_run('\n分析师：析数数（数据分析专家）').italic = True
p.add_run('\n数据来源：MEMORY.md 44场复盘记录').italic = True

# 保存文件
output_path = r'C:\Users\84019\Desktop\足球博彩Rule系统优化分析报告.docx'
doc.save(output_path)
print(f'Word文档已创建：{output_path}')
